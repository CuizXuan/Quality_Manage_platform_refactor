"""
文档构建器 — 规则驱动的文档生成
从 doc-generator 迁入并修复 save() 方法
"""

import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from .datamodel import DocumentTree, LeafData


class DocBuilder:
    """文档构建器"""

    def __init__(self, template_path: str, rule: dict):
        self.rule = rule
        self._fmt = rule.get('formatting', {})
        self._level_rules = rule.get('levelRules', {})
        self._blocks = rule.get('leafBlocks', [])
        self._external_data: dict[str, LeafData] = {}
        self._doc = Document(template_path)
        self._clear_template()
        self._table_style = None
        try:
            self._table_style = self._doc.styles.get(self._fmt.get('tableStyle', ''))
        except Exception:
            pass

    def _clear_template(self):
        """清空模板内容，保留页面设置"""
        body = self._doc.element.body
        for child in list(body):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag != 'sectPr':
                body.remove(child)

    def _set_font(self, run):
        try:
            run.font.name = self._fmt.get('bodyFont', '仿宋_GB2312')
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self._fmt.get('bodyFont', '仿宋_GB2312'))
            run.font.size = Pt(self._fmt.get('bodySize', 14))
        except Exception:
            pass

    def _add_body(self, text=""):
        p = self._doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Pt(self._fmt.get('bodyIndent', 28))
        for r in p.runs:
            self._set_font(r)
        return p

    def _add_table(self, flow_rows):
        if not flow_rows:
            self._add_body('')
            return
        t = self._doc.add_table(rows=len(flow_rows), cols=2)
        t.style = 'Table Grid'
        for i, row in enumerate(flow_rows):
            t.rows[i].cells[0].text = row.action
            t.rows[i].cells[1].text = row.response
            for cell in t.rows[i].cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Pt(0)
                    if self._table_style:
                        p.style = self._table_style
                    else:
                        for r in p.runs:
                            self._set_font(r)
        return t

    def _get_content_level(self, leaf_h_level: int) -> int:
        default = self._level_rules.get('defaultContentLevel', 6)
        overrides = self._level_rules.get('overrides', {})
        return overrides.get(str(leaf_h_level), default)

    def set_external_data(self, data: dict[str, LeafData]):
        self._external_data.update(data)
        return self

    def build(self, tree: DocumentTree):
        doc = self._doc
        doc.add_heading('定制软件', 2)

        for plat in tree.platforms:
            for sec in plat.sections:
                doc.add_heading(sec.name, 3)
                for leaf in sec.direct_leaves:
                    doc.add_heading(leaf.name, 4)
                    self._build_leaf_blocks(leaf)

                for grp in sec.groups:
                    if grp.is_leaf:
                        ref = getattr(grp, '_leaf_ref', None)
                        doc.add_heading(grp.name, 4)
                        if ref:
                            self._build_leaf_blocks(ref)
                        else:
                            self._build_leaf_blocks_for_name(grp.name, grp.heading_level)
                    else:
                        doc.add_heading(grp.name, 4)
                        for leaf in grp.children:
                            doc.add_heading(leaf.name, 5)
                            self._build_leaf_blocks(leaf)
        return self

    def _build_leaf_blocks(self, leaf):
        self._build_leaf_blocks_for_name(leaf.name, leaf.heading_level, leaf.data)

    def _build_leaf_blocks_for_name(self, name: str, heading_level: int, data=None):
        cl = self._get_content_level(heading_level)

        for block in self._blocks:
            title = block['title']
            source = block['source']
            self._doc.add_heading(title, cl)

            if source.startswith('db_') or source.startswith('api_'):
                ext_data = self._external_data.get(name)
                if ext_data:
                    if source in ('db_desc', 'api_summary'):
                        self._add_body(ext_data.desc or '暂无')
                    elif source in ('db_columns', 'api_endpoints', 'api_params'):
                        if ext_data.flow:
                            self._add_table(ext_data.flow)
                        else:
                            self._add_body('')
                    elif source in ('db_indexes', 'api_responses'):
                        self._add_body(ext_data.desc or '暂无')
                    else:
                        self._add_body('')
                else:
                    self._add_body('暂无')
            elif source == 'desc' and data:
                self._add_body(data.desc or '暂无')
            elif source == 'flow' and data:
                if data.flow:
                    self._add_table(data.flow)
                else:
                    self._add_body('')
            else:
                self._add_body('')

    def save(self, output_path: str):
        """保存文档到指定路径（Windows/Linux 兼容）"""
        self._doc.save(output_path)
        return self

    @property
    def paragraph_count(self) -> int:
        return len(self._doc.paragraphs)

    @property
    def table_count(self) -> int:
        return len(self._doc.tables)